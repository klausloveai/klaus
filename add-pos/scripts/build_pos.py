#!/usr/bin/env python3
"""
build_pos.py — Fill the firm POS template, append a copy to the end of each document
being served, and generate a copy-paste service email, for 凌图律所 / Lingtu Law
(Law Office of Shenqi Cai APC).

Usage:
    python3 build_pos.py config.json

The firm POS template ("POS Template.docx", Drive id 19BhkRUm99mGnajKmAP-vaQCoFzfZnWCU)
is a pleading-format Proof of Service with double side-rules, line numbers, a Word DATE
auto-field (fills the serve day), a footer PAGE field + title, and {{tokens}}. This script:
  1. loads the template,
  2. fills every token with the case/serve info,
  3. strips any yellow highlight from the filled text (clean served output),
  4. adds a ~0.13" right-margin gap so a long documents-served list never touches the rule,
  5. for EACH document: sets the footer title to that document's own caption title and
     sets the POS page-numbering to continue the host document (start = host pages + 1),
  6. converts the filled POS to PDF (LibreOffice) and appends it after the document's
     ORIGINAL pages (attorney pages are never modified),
  7. writes each combined PDF named "<original filename> (with POS).pdf",
  8. generates a copy-paste service email (prints it and writes email.txt to outdir).

It DRAFTS ONLY — it never sends, serves, saves a Gmail draft, or applies a signature.
The declarant signs the POS and the user serves.

Requires: python-docx, pypdf, LibreOffice (`soffice`) on PATH.

config.json shape (see also references/):
{
  "template": "/abs/POS Template.docx",
  "outdir": "/Users/klaus/Downloads",
  "skip_build": false,                       // true = skip appending POS, only make the email
  "documents_served_lines": [                // single source of truth for both POS + email
    "Claimant's Objection to Respondent's Notice of Taking Deposition of Zhiping Liu",
    "Claimant's Request for Production of Documents to Respondent (Set No. One)"
  ],
  "set_no": "One",
  "case_name": "Zhiping Liu v. State Farm Mutual Automobile Insurance Company",
  "court_case_no": "In the Matter of the Arbitration Between ... — Claim No. 75-78X9-98Q",
  "counsel_name_sbn": "JOHN P. YASUDA, ESQ., SBN 133025",
  "firm": "JAMES T. SHOTT & ASSOCIATES",
  "street": "611 Anton Boulevard, Suite 900",
  "city_zip": "Costa Mesa, California 92626",
  "telephone": "(714) 435-7511",
  "fax": "(855) 396-4486",
  "eservice_email": "Cali-Law-Costa-Mesa-Clc@StateFarm.com",   // DESIGNATED service address ONLY
  "party_role": "Respondent",
  "party_name": "STATE FARM MUTUAL AUTOMOBILE INSURANCE COMPANY",
  "declarant": "Klaus Liu",
  "documents": [                             // "out" auto-derived from src basename
    {"src": "/abs/Zhiping Liu - Objection to Deposition.pdf",
     "footer_title": "CLAIMANT'S OBJECTION TO RESPONDENT'S NOTICE OF TAKING DEPOSITION OF ZHIPING LIU"}
  ],
  "email": {
    "claimant": "Zhiping Liu",
    "matter_no": "Claim No. 75-78X9-98Q",
    "service_desc": "Claimant's Discovery Requests and Objection to Deposition",
    "designated_letter_date": "June 19, 2026",
    "to": "Cali-Law-Costa-Mesa-Clc@StateFarm.com",             // default = eservice_email
    "cc": ["john.yasuda@statefarm.com", "hernan.s@lingtulaw.com", "klaus@lingtulaw.com"]
  }
}
"""
import json, sys, os, subprocess
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pypdf

RIGHT_GAP = Inches(0.13)


def inline_list(lines):
    """['A','B','C'] -> '1. A; 2. B; and 3. C'  (single item -> '1. A')."""
    if not lines:
        return ""
    if len(lines) == 1:
        return f"1. {lines[0]}"
    parts = [f"{i+1}. {t}" for i, t in enumerate(lines[:-1])]
    parts.append(f"and {len(lines)}. {lines[-1]}")
    return "; ".join(parts)


def documents_served_string(cfg):
    if cfg.get("documents_served"):          # explicit inline string wins if provided
        return cfg["documents_served"]
    return inline_list(cfg.get("documents_served_lines", []))


def para_replace(p, old, new):
    runs = p.runs
    texts = [r.text for r in runs]
    full = "".join(texts)
    i = full.find(old)
    if i < 0:
        return False
    s, e = i, i + len(old)
    pos = 0
    first = True
    for r, t in zip(runs, texts):
        rs, re_ = pos, pos + len(t)
        pos = re_
        if re_ <= s or rs >= e:
            continue
        ls = max(s, rs) - rs
        le = min(e, re_) - rs
        if first:
            r.text = t[:ls] + new + t[le:]
            first = False
        else:
            r.text = t[:ls] + t[le:]
    return True


def strip_highlight(p):
    for r in p.runs:
        try:
            r.font.highlight_color = None
        except Exception:
            pass


def all_body_and_table_paras(d):
    ps = list(d.paragraphs)
    for t in d.tables:
        for row in t.rows:
            for c in row.cells:
                ps += list(c.paragraphs)
    return ps


def set_pgnum_start(sec, start):
    pg = sec._sectPr.find(qn("w:pgNumType"))
    if pg is None:
        pg = OxmlElement("w:pgNumType")
        sec._sectPr.append(pg)
    pg.set(qn("w:start"), str(start))


def out_name(src):
    base = os.path.splitext(os.path.basename(src))[0]
    return f"{base} (with POS).pdf"


def build_one(cfg, doc, served):
    d = Document(cfg["template"])
    tokens = {
        "{{CASE NAME}}": cfg["case_name"],
        "{{COURT & CASE NO. (or: In the Matter of the Arbitration... Claim No. ___)}}": cfg["court_case_no"],
        "{{OPPOSING COUNSEL NAME, SBN}}": cfg["counsel_name_sbn"],
        "{{FIRM NAME}}": cfg["firm"],
        "{{STREET ADDRESS}}": cfg["street"],
        "{{CITY, STATE ZIP}}": cfg["city_zip"],
        "{{TELEPHONE}}": cfg["telephone"],
        "{{FAX}}": cfg["fax"],
        "{{E-SERVICE EMAIL}}": cfg["eservice_email"],
        "{{PARTY ROLE (Defendant/Respondent)}}": cfg["party_role"],
        "{{PARTY NAME}}": cfg["party_name"],
        "{{DECLARANT NAME}}": cfg["declarant"],
    }
    for p in all_body_and_table_paras(d):
        para_replace(p, "{{TITLE(S) OF DOCUMENT(S) SERVED}}, {{SET NO.}}", served)
        para_replace(p, "{{TITLE(S) OF DOCUMENT(S) SERVED}}", served)
        para_replace(p, "{{SET NO.}}", cfg.get("set_no", "One"))
        for k, v in tokens.items():
            para_replace(p, k, v)
        strip_highlight(p)
    for p in d.paragraphs:
        p.paragraph_format.right_indent = RIGHT_GAP
    n = len(pypdf.PdfReader(doc["src"]).pages)
    start = n + 1
    for sec in d.sections:
        for p in sec.footer.paragraphs:
            para_replace(p, "{{TITLE(S) OF DOCUMENT(S) SERVED}} [{{SET NO.}}]", doc["footer_title"])
            para_replace(p, "{{TITLE(S) OF DOCUMENT(S) SERVED}}", doc["footer_title"])
            para_replace(p, "[{{SET NO.}}]", "")
            para_replace(p, "{{SET NO.}}", "")
            strip_highlight(p)
        set_pgnum_start(sec, start)

    tmp_docx = os.path.join(cfg["outdir"], "_pos_tmp.docx")
    d.save(tmp_docx)
    subprocess.run(["soffice", "--headless", "--convert-to", "pdf", "--outdir", cfg["outdir"], tmp_docx],
                   capture_output=True)
    tmp_pdf = tmp_docx[:-5] + ".pdf"
    left = [p.text for p in Document(tmp_docx).paragraphs if "{{" in p.text]
    w = pypdf.PdfWriter()
    for pg in pypdf.PdfReader(doc["src"]).pages:
        w.add_page(pg)
    pos_pages = pypdf.PdfReader(tmp_pdf).pages
    for pg in pos_pages:
        w.add_page(pg)
    out_path = os.path.join(cfg["outdir"], doc.get("out") or out_name(doc["src"]))
    with open(out_path, "wb") as fh:
        w.write(fh)
    for f in (tmp_docx, tmp_pdf):
        try:
            os.remove(f)
        except OSError:
            pass
    return out_path, n, len(pos_pages), start, left


def make_email(cfg):
    e = cfg.get("email", {})
    lines = cfg.get("documents_served_lines", [])
    # To = the designated service address(es). Per James Zhan's rule, in a multi-defendant
    # case serve TO ALL defense counsel (all counsel have a right to know) — pass a list.
    to = e.get("to") or cfg.get("eservice_email", "[DESIGNATED SERVICE ADDRESS]")
    if isinstance(to, (list, tuple)):
        to = "; ".join(to)
    cc = e.get("cc", [])
    case_name = e.get("case_name") or cfg.get("case_name", "")
    desc = e.get("service_desc", "the following documents")
    claimant = e.get("claimant", "our client")
    ltr_date = e.get("designated_letter_date", "")
    desig_clause = (f", and to the electronic service address designated in your office's "
                    f"{ltr_date} correspondence," if ltr_date else ",")
    # For insurance carriers, put the Claim No. + Defense File No. in the subject and body.
    claim_no = e.get("claim_no", "")
    defense_file_no = e.get("defense_file_no", "")
    ids = []
    if claim_no:
        ids.append(f"Claim No.: {claim_no}")
    if defense_file_no:
        ids.append(f"Defense File No.: {defense_file_no}")
    if ids:
        subject = f"{case_name} | {' | '.join(ids)} — Service of {desc}"
        ref = [b.replace(":", "") for b in ids]
        matter_ref = f" ({'; '.join(ref)})"
    else:
        matter_no = e.get("matter_no") or cfg.get("court_case_no", "")
        subject = f"{case_name} / Service of {desc}"
        matter_ref = f" ({matter_no})" if matter_no else ""
    # numbered body list
    body_items = []
    for i, t in enumerate(lines):
        sep = ";" if i < len(lines) - 1 else ""
        if i == len(lines) - 1 and len(lines) > 1:
            body_items[-1] = body_items[-1][:-1] + "; and"
        body_items.append(f"{i+1}. {t}{sep}")
    if body_items:
        body_items[-1] = body_items[-1].rstrip(";") + "."
    doc_block = "\n".join(body_items)

    # Default body = simple/conversational (the carrier IDs live in the subject; the POS
    # attached to each document does the formal §1010.6 work). desig_clause/claimant/
    # matter_ref are computed above and available if a longer recital is ever wanted.
    _ = (desig_clause, claimant, matter_ref)
    body = f"""Dear Counsel,

Please see attachments.

{doc_block}

A Proof of Service is included with each document. Kindly confirm receipt at your convenience.

Thank you,

Klaus Liu | Director of Case Management
Lingtu Law Office
13191 Crossroads Pkwy N, Suite 295, City of Industry, CA 91746
Direct: (626) 479-2207 | Fax: (626) 479-2207
"""
    display = f"To: {to}\nCc: {', '.join(cc)}\nSubject: {subject}\n\n{body}"
    return {"to": to, "cc": cc, "subject": subject, "body": body, "display": display}


def create_gmail_draft(cfg, email, attachments):
    """Create a Gmail DRAFT (klaus@) with the service email + POS'd PDFs attached.
    DRAFT ONLY — never sends. Returns the draft id (or None)."""
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText
    from email.mime.application import MIMEApplication
    msg = MIMEMultipart()
    msg["From"] = cfg.get("email", {}).get("from", "klaus@lingtulaw.com")
    msg["To"] = email["to"]
    if email["cc"]:
        msg["Cc"] = ", ".join(email["cc"])
    msg["Subject"] = email["subject"]
    msg.attach(MIMEText(email["body"], "plain", "utf-8"))
    for f in attachments:
        with open(f, "rb") as fh:
            part = MIMEApplication(fh.read(), _subtype="pdf")
        part.add_header("Content-Disposition", "attachment", filename=os.path.basename(f))
        msg.attach(part)
    eml = "_draft.eml"                      # gws --upload requires the file inside its cwd
    with open(os.path.join(cfg["outdir"], eml), "wb") as fh:
        fh.write(msg.as_bytes())
    r = subprocess.run(
        ["gws", "gmail", "users", "drafts", "create", "--params", '{"userId":"me"}',
         "--upload", eml, "--upload-content-type", "message/rfc822"],
        cwd=cfg["outdir"], capture_output=True, text=True)
    try:
        os.remove(os.path.join(cfg["outdir"], eml))
    except OSError:
        pass
    out = "\n".join(l for l in r.stdout.splitlines() if "keyring" not in l)
    try:
        return json.loads(out).get("id")
    except Exception:
        return None


def main():
    cfg = json.load(open(sys.argv[1]))
    os.makedirs(cfg["outdir"], exist_ok=True)
    served = documents_served_string(cfg)

    if not cfg.get("skip_build"):
        print(f"Template: {cfg['template']}")
        print(f"Service address (POS): {cfg['eservice_email']}")
        for doc in cfg["documents"]:
            out, n, pn, start, left = build_one(cfg, doc, served)
            flag = "OK" if not left else f"LEFTOVER TOKENS: {left}"
            print(f"  {os.path.basename(out)}: {n} doc + {pn} POS = {n+pn}pp | POS pages {start}-{start+pn-1} | {flag}")
    else:
        print("(skip_build) — email only")

    email = make_email(cfg)
    with open(os.path.join(cfg["outdir"], "service_email.txt"), "w") as fh:
        fh.write(email)
    print("\n================= COPY-PASTE SERVICE EMAIL =================\n")
    print(email)
    print("===========================================================")


if __name__ == "__main__":
    main()
