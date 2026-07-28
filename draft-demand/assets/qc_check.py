#!/usr/bin/env python3
"""
QC checker for a Lingtu Law demand letter (.docx).

Catches the errors that must NEVER reach a carrier: leftover template artifacts,
wrong math, internally inconsistent totals, wrong-carrier case cites, a deadline
that isn't a Monday, and broken exhibit references. It also PRINTS the source-
derived values (ICD codes, MRI findings, provider rows, per-diem dates, claim
data) so the drafter can eyeball them against the actual records/bills — the
"don't fabricate / don't miss a number" check a script can't make on its own.

Usage:  python3 qc_check.py "/path/to/<Client> - 3P Demand Letter.docx"
Exit code 0 = no hard FAILs (WARN/INFO may remain). Exit code 1 = at least one FAIL.
"""
import sys, re, datetime
import docx

FAILS, WARNS, INFOS = [], [], []
def fail(m): FAILS.append(m)
def warn(m): WARNS.append(m)
def info(m): INFOS.append(m)

def money(s):
    """Parse '$9,510.50' / '16.0' / '365' -> float, else None."""
    if s is None: return None
    t = re.sub(r'[^0-9.\-]', '', str(s))
    if t in ('', '.', '-'): return None
    try: return float(t)
    except ValueError: return None

def approx(a, b, tol=0.01):
    return a is not None and b is not None and abs(a - b) <= tol

def cellrows(tb):
    return [[c.text.strip().replace('\n', ' ') for c in r.cells] for r in tb.rows]

def find_table(doc, needle):
    for tb in doc.tables:
        for r in cellrows(tb):
            if any(needle.lower() in c.lower() for c in r):
                return tb
    return None

# ---------------------------------------------------------------- load
if len(sys.argv) < 2:
    print("usage: qc_check.py <demand.docx>"); sys.exit(2)
path = sys.argv[1]
doc = docx.Document(path)

# collect all paragraph runs (body + tables) for artifact scans
def iter_runs(d):
    for p in d.paragraphs:
        for r in p.runs: yield r
    for tb in d.tables:
        for row in tb.rows:
            for c in row.cells:
                for p in c.paragraphs:
                    for r in p.runs: yield r
ALLTEXT = "\n".join(p.text for p in doc.paragraphs)
for tb in doc.tables:
    for r in cellrows(tb):
        ALLTEXT += "\n" + " | ".join(r)

# ---------------------------------------------------------------- 1. leftover artifacts
if '{{' in ALLTEXT or '}}' in ALLTEXT:
    toks = sorted(set(re.findall(r'\{\{[^}]{0,60}', ALLTEXT)))
    fail(f"Leftover template token(s): {toks[:8]}")
red = [r.text.strip() for r in iter_runs(doc)
       if r.font.color and r.font.color.rgb and str(r.font.color.rgb) == 'C00000' and r.text.strip()]
if red:
    fail(f"Leftover red instruction text ({len(red)}): {red[0][:70]!r} ...")
yel = [r.text.strip() for r in iter_runs(doc)
       if r.font.highlight_color and str(r.font.highlight_color) not in ('None', 'AUTO (0)') and r.text.strip()]
if yel:
    fail(f"Leftover yellow-highlighted text ({len(yel)}): {yel[0][:50]!r} ...")
for ph in ['[describe', '[insert', '[replace', 'TODO', 'XXXX', 'lashinelaw', 'Ms. Faye']:
    if ph.lower() in ALLTEXT.lower():
        fail(f"Placeholder / wrong-firm artifact present: {ph!r}")

# ---------------------------------------------------------------- 2. citation accuracy
for bad in ['Murphy v. Farmers', 'Betts v. Farmers', 'Cain v. Farmers', 'Cal.App3d',
            'Commercial Union Insurance Company v. Safeway']:
    if bad in ALLTEXT:
        fail(f"Known-bad citation string: {bad!r}")

# ---------------------------------------------------------------- 3. gender consistency
neutral = re.sub(r'\b(he or she|his or her|him or her|he/she|his/her)\b', '', ALLTEXT, flags=re.I)
masc = len(re.findall(r'\b(he|his|him|himself|Mr\.)\b', neutral))
fem  = len(re.findall(r'\b(she|her|hers|herself|Ms\.)\b', neutral))
if masc and fem:
    warn(f"Mixed gender pronouns (masc={masc}, fem={fem}) — confirm a swap wasn't missed.")

# ---------------------------------------------------------------- 4. Past-Medical table math
pm = find_table(doc, 'Amount Charged')
if pm:
    rows = cellrows(pm)
    hdr = rows[0]
    try:
        ci_amt = next(i for i, h in enumerate(hdr) if 'amount' in h.lower())
        ci_vis = next(i for i, h in enumerate(hdr) if 'visit' in h.lower())
        ci_ref = next((i for i, h in enumerate(hdr) if 'reference' in h.lower()), None)
    except StopIteration:
        ci_amt = ci_vis = ci_ref = None
    if ci_amt is not None:
        data, total_amt, total_vis = [], None, None
        for r in rows[1:]:
            amt = money(r[ci_amt]) if ci_amt < len(r) else None
            prov = r[0].strip()
            if prov == '' and amt is not None:          # the black total row
                total_amt = amt
                total_vis = money(r[ci_vis]) if ci_vis < len(r) else None
            elif amt is not None:
                data.append((prov, amt, money(r[ci_vis]) if ci_vis is not None and ci_vis < len(r) else None,
                             r[ci_ref] if ci_ref is not None and ci_ref < len(r) else ''))
        s_amt = sum(d[1] for d in data)
        if total_amt is not None and not approx(s_amt, total_amt):
            fail(f"Past-Medical $ total mismatch: rows sum ${s_amt:,.2f} != table total ${total_amt:,.2f}")
        s_vis = sum(d[2] for d in data if d[2] is not None)
        if total_vis is not None and not approx(s_vis, total_vis):
            fail(f"Past-Medical visit total mismatch: rows sum {s_vis:g} != table total {total_vis:g}")
        for prov, amt, vis, ref in data:
            info(f"PROVIDER (verify vs bill): {prov} | visits={vis:g} | ${amt:,.2f} | {ref}")
        PAST_MED = total_amt if total_amt is not None else s_amt
    else:
        PAST_MED = None; warn("Past-Medical table found but columns unreadable.")
else:
    PAST_MED = None; warn("No Past-Medical table found.")

# ---------------------------------------------------------------- 5. Future-Medical table math
fm = find_table(doc, 'Per Year') or find_table(doc, 'Procedure')
FUT_MED = None
if fm:
    rows = cellrows(fm)
    costs, total = [], None
    for r in rows[1:]:
        last = money(r[-1])
        if r[0].strip().lower() == 'total' or all(c.strip().lower() == 'total' for c in r[:-1] if c.strip()):
            total = last
        elif last is not None:
            costs.append(last)
    if costs:
        s = sum(costs)
        if total is not None and not approx(s, total):
            fail(f"Future-Medical total mismatch: rows sum ${s:,.2f} != ${total:,.2f}")
        FUT_MED = total if total is not None else s
        if not approx(FUT_MED, 3920.00):
            info(f"Future-Medical = ${FUT_MED:,.2f} (default conservative is $3,920; higher only if a PM/ortho report recommends injections/surgery — confirm a report is cited).")

# ---------------------------------------------------------------- 6. Pain & Suffering per-diem math
ps = find_table(doc, 'Total Pain and Suffering')
PS_TOTAL = None
if ps:
    sec = None
    vals = {'initial': {}, 'subsequent': {}}
    grand = None
    for r in cellrows(ps):
        label = r[0].strip(); val = money(r[-1]) if len(r) > 1 else None
        ll = label.lower()
        if 'initial pain' in ll and r[0] == r[-1]: sec = 'initial'; continue
        if 'subsequent pain' in ll and r[0] == r[-1]: sec = 'subsequent'; continue
        if ll.startswith('total pain and suffering'): grand = val; continue
        if sec:
            if 'waking hours' in ll:          vals[sec]['hours'] = val   # must precede 'day' (label = "Waking Hours/Day")
            elif 'compensation/hour' in ll:   vals[sec]['rate'] = val
            elif 'day' in ll:                 vals[sec]['days'] = val
            elif ll.startswith('total'):      vals[sec]['total'] = val
    sub_sum = 0
    for sec in ('initial', 'subsequent'):
        v = vals[sec]
        if all(k in v and v[k] is not None for k in ('days', 'hours', 'rate', 'total')):
            calc = v['days'] * v['hours'] * v['rate']
            if not approx(calc, v['total']):
                fail(f"Per-diem {sec}: {v['days']:g}×{v['hours']:g}×${v['rate']:g} = ${calc:,.2f} but table shows ${v['total']:,.2f}")
            sub_sum += v['total']
            info(f"PER-DIEM {sec}: {v['days']:g} days × {v['hours']:g} hrs × ${v['rate']:g}/hr = ${v['total']:,.2f}")
        else:
            warn(f"Per-diem {sec} section incomplete: {v}")
    if grand is not None:
        if not approx(sub_sum, grand):
            fail(f"P&S grand total mismatch: initial+subsequent ${sub_sum:,.2f} != ${grand:,.2f}")
        PS_TOTAL = grand

# ---------------------------------------------------------------- 7. Damages summary + cross-table
ds = find_table(doc, 'Total Damages')
if ds:
    parts, total = {}, None
    for r in cellrows(ds):
        label = r[0].strip(); val = money(r[-1]) if len(r) > 1 else None
        ll = label.lower()
        if r[0] == r[-1]: continue
        if 'total damages' in ll: total = val
        elif 'past medical' in ll: parts['past'] = val
        elif 'future medical' in ll: parts['future'] = val
        elif 'loss of income' in ll or 'lost income' in ll or 'wage' in ll: parts['loi'] = val
        elif 'pain and suffering' in ll: parts['ps'] = val
    s = sum(v for v in parts.values() if v is not None)
    if total is not None and not approx(s, total):
        fail(f"Total Damages mismatch: components sum ${s:,.2f} != ${total:,.2f}  (parts={ {k: f'${v:,.2f}' for k,v in parts.items()} })")
    # cross-check against the source tables
    if PAST_MED is not None and 'past' in parts and not approx(PAST_MED, parts['past']):
        fail(f"Past-Medical summary ${parts['past']:,.2f} != Past-Medical table ${PAST_MED:,.2f}")
    if FUT_MED is not None and 'future' in parts and not approx(FUT_MED, parts['future']):
        fail(f"Future-Medical summary ${parts['future']:,.2f} != Future-Medical table ${FUT_MED:,.2f}")
    if PS_TOTAL is not None and 'ps' in parts and not approx(PS_TOTAL, parts['ps']):
        fail(f"P&S summary ${parts['ps']:,.2f} != P&S table ${PS_TOTAL:,.2f}")
    if total is not None: info(f"TOTAL DAMAGES = ${total:,.2f}")

# ---------------------------------------------------------------- 8. deadline: weekday label must be correct + a Monday
MONTHS = 'January|February|March|April|May|June|July|August|September|October|November|December'
dm = re.search(r'(Monday|Tuesday|Wednesday|Thursday|Friday|Saturday|Sunday),?\s+(' + MONTHS + r')\s+(\d{1,2}),?\s+(\d{4})', ALLTEXT)
if not dm:  # fall back to "remain open until <date>" / "by <date>"
    dm2 = re.search(r'(?:until|by|before|no later than)\s+(?:\w+,?\s+)?(' + MONTHS + r')\s+(\d{1,2}),?\s+(\d{4})', ALLTEXT)
    if dm2:
        try:
            dt = datetime.datetime.strptime(f"{dm2.group(1)} {dm2.group(2)} {dm2.group(3)}", "%B %d %Y").date()
            wd = dt.strftime('%A')
            if wd != 'Monday': warn(f"Deadline {dt.isoformat()} is a {wd}, not a Monday (firm convention = a Monday ~30 days out).")
            info(f"DEADLINE: {dt.isoformat()} ({wd})")
        except ValueError: pass
    else:
        warn("No demand deadline date found (expected a weekday-labeled date or 'remain open until <date>').")
else:
    stated_wd = dm.group(1)
    try:
        dt = datetime.datetime.strptime(f"{dm.group(2)} {dm.group(3)} {dm.group(4)}", "%B %d %Y").date()
        real_wd = dt.strftime('%A')
        if stated_wd != real_wd:
            fail(f"Deadline weekday WRONG: letter says {stated_wd}, {dt.isoformat()} but that date is a {real_wd}.")
        elif real_wd != 'Monday':
            warn(f"Deadline {dt.isoformat()} is a {real_wd}, not a Monday (firm convention = a Monday ~30 days out).")
        info(f"DEADLINE: {stated_wd}, {dt.isoformat()} ({real_wd})")
    except ValueError:
        warn(f"Could not parse deadline date {dm.group(0)!r}.")

# ---------------------------------------------------------------- 9. exhibit references resolve
ex = find_table(doc, 'Exhibit No')
if ex:
    listed = set(int(x) for x in re.findall(r'Exhibit\s+(\d+)', "\n".join(" ".join(r) for r in cellrows(ex))))
    refd = set(int(x) for x in re.findall(r'Exhibit\s+(\d+)', ALLTEXT))
    missing = refd - listed
    if missing:
        fail(f"Body references Exhibit(s) {sorted(missing)} not in the exhibit list {sorted(listed)}")
    info(f"EXHIBITS listed: {sorted(listed)}")

# ---------------------------------------------------------------- 10. required fields present
icd = find_table(doc, 'ICD Code')
if icd:
    codes = [r[0] for r in cellrows(icd)[1:] if r[0].strip()]
    info(f"ICD CODES ({len(codes)}) — verify EACH appears in the records: {', '.join(codes)}")
for mri in re.findall(r'(At [CTLS]\d.*?thecal sac\.|At [CTLS]\d.*?subarachnoid space.*?\.)', ALLTEXT, re.S):
    info("MRI FINDING (verify verbatim vs radiology report): " + re.sub(r'\s+', ' ', mri)[:160])
for label, pat in [('Claim No.', r'Claim No\.?\s*[:#]?\s*([A-Za-z0-9\-]+)'),
                   ('Handler phone', r'\((\d{3})\)\s*\d{3}-\d{4}')]:
    if not re.search(pat, ALLTEXT):
        warn(f"Could not find {label} in the letter — confirm it's present.")

# ---------------------------------------------------------------- report
print("=" * 72)
print(f"QC REPORT — {path.split('/')[-1]}")
print("=" * 72)
if FAILS:
    print(f"\n❌ FAIL ({len(FAILS)}) — MUST fix before sending:")
    for x in FAILS: print("   • " + x)
if WARNS:
    print(f"\n⚠️  WARN ({len(WARNS)}) — confirm each:")
    for x in WARNS: print("   • " + x)
print(f"\nℹ️  VERIFY-AGAINST-SOURCE ({len(INFOS)}) — eyeball every number/fact vs the records & intake:")
for x in INFOS: print("   • " + x)
print("\n" + ("RESULT: ❌ NOT CLEAN — fix all FAILs, then re-run." if FAILS else
             "RESULT: ✅ no hard failures. Still confirm every WARN and VERIFY item against the source."))
sys.exit(1 if FAILS else 0)
